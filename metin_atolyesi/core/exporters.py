from __future__ import annotations

from pathlib import Path

from .models import Project


def export_txt(project: Project, path: Path) -> None:
    content = "\n\n".join(f"[{p.label or p.page_index + 1}]\n{p.text}" for p in project.pages)
    path.write_text(content, encoding="utf-8")


def export_excel(project: Project, path: Path) -> None:
    from openpyxl import Workbook

    wb = Workbook()
    ws = wb.active
    ws.title = "Dizin"
    headers = ["Madde basi", "Koken", "Anlam", "Kullanim", "Ek", "Sayfa/Varak", "Not", "Guven", "Goruntu"]
    ws.append(headers)
    for item in project.vocabulary:
        ws.append([
            item.headword,
            item.origin,
            item.meaning,
            item.usage,
            item.suffixes,
            item.location,
            item.note,
            item.confidence,
            item.image_path,
        ])
    text_ws = wb.create_sheet("Metin")
    text_ws.append(["Sayfa", "Metin"])
    for page in project.pages:
        text_ws.append([page.label or page.page_index + 1, page.text])
    wb.save(path)


def export_word(project: Project, path: Path) -> None:
    from docx import Document

    doc = Document()
    doc.add_heading(project.name, level=1)
    for page in project.pages:
        doc.add_heading(str(page.label or page.page_index + 1), level=2)
        doc.add_paragraph(page.text)
    if project.vocabulary:
        doc.add_heading("Dizin / Soz Varligi", level=2)
        table = doc.add_table(rows=1, cols=6)
        hdr = table.rows[0].cells
        for cell, label in zip(hdr, ["Madde", "Koken", "Anlam", "Kullanim", "Ek", "Yer"]):
            cell.text = label
        for item in project.vocabulary:
            row = table.add_row().cells
            values = [item.headword, item.origin, item.meaning, item.usage, item.suffixes, item.location]
            for cell, value in zip(row, values):
                cell.text = str(value)
    doc.save(path)
